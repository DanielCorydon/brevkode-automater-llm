import tempfile
import os
import traceback
from docx import Document
from components.tools import convert_text_to_mergefields, text_to_word_docx

try:
    print("Starting test...")
    
    # Create a temporary file with our test content
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        test_file = f.name
        print(f"Created temporary file: {test_file}")
    
    # Create a Word document with our test string
    test_text = '{ IF "J" "{ MERGEFIELD ab-borger-enlig-ved-aeldrecheck-berettigelse }" "dine" "din og din samlever/ægtefælles" }'
    print(f"Test text: {test_text}")
    text_to_word_docx(test_text, test_file)
    print("Created Word document from text")

    # Try to convert the text to actual mergefields
    print("Attempting to convert text to mergefields...")
    success, debug_info = convert_text_to_mergefields(test_file, test_file)
    print(f'Success: {success}')
    print(f'Debug info: {debug_info}')
    
except Exception as e:
    print(f"ERROR: {str(e)}")
    print(traceback.format_exc())

# Now try to open the file to verify it worked
try:
    doc = Document(test_file)
    print(f"Document has {len(doc.paragraphs)} paragraphs")
    print(f"First paragraph contains {len(doc.paragraphs[0].runs)} runs")
except Exception as e:
    print(f"Error opening document: {str(e)}")
