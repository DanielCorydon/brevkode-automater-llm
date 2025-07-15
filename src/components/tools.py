# --- Tool: Create a Word document from text ---
from docx import Document
import io


def text_to_word_docx(text: str, output_path: str = None) -> bytes:
    """
    Create a Word (.docx) document from the given text. Each line becomes a paragraph.
    If output_path is provided, saves the file there. Returns the bytes of the document.

    Args:
        text (str): The input text.
        output_path (str, optional): Path to save the .docx file. If None, only returns bytes.

    Returns:
        bytes: The .docx file as bytes.
    """
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()
    if output_path:
        with open(output_path, "wb") as f:
            f.write(doc_bytes)
    return doc_bytes


"""
Tool functions for document processing and mapping replacements.
"""

import os
import pandas as pd


# --- Minimal Excel mapping loader ---
def load_excel_mapping(path):
    try:
        df = pd.read_excel(path, sheet_name=None)
        # Try to find the right sheet ("query" or first sheet)
        sheet = df["query"] if "query" in df else next(iter(df.values()))
        # Expect columns: "Titel" and "Nøgle"
        if "Titel" in sheet.columns and "Nøgle" in sheet.columns:
            return dict(zip(sheet["Titel"], sheet["Nøgle"]))
    except Exception:
        pass
    return {}


# Path to default mapping Excel (relative to this file)
DEFAULT_MAPPING_PATH = os.path.join(
    os.path.dirname(os.path.dirname(__file__)),
    "documents",
    "Liste over alle nøgler.xlsx",
)
MAPPINGS_DICT = load_excel_mapping(DEFAULT_MAPPING_PATH)


def search_and_replace(text: str, search: str, replace: str) -> str:
    """
    Search for all occurrences of 'search' in 'text' and replace them with 'replace'.

    Args:
        text (str): The input text.
        search (str): The substring to search for.
        replace (str): The substring to replace with.

    Returns:
        str: The modified text with replacements.
    """
    return text.replace(search, replace)


# --- New tool: Replace any occurrence of a Title (Titel) with its corresponding Nøgle ---
def replace_titels_with_nogle(text: str, replacement_template: str) -> str:
    """
    Replace all occurrences of each 'Titel' in the mapping Excel with the replacement_template.
    If '<NØGLE>' is present in the template, it is replaced with the corresponding key (Nøgle).

    Args:
        text (str): The input text.
        replacement_template (str): The template to replace each Titel. Use '<NØGLE>' to insert the key.

    Returns:
        str: The modified text.
    """
    result = text
    for titel, nøgle in MAPPINGS_DICT.items():
        if titel:
            if "<NØGLE>" in replacement_template:
                replacement = replacement_template.replace("<NØGLE>", str(nøgle))
            else:
                replacement = replacement_template
            result = result.replace(str(titel), replacement)
    return result


# --- Tool: Convert mergefield-like text and IF fields to actual Word fields ---
from docx.oxml import OxmlElement
import re


def convert_text_to_mergefields(docx_path: str, output_path: str = None) -> tuple:
    """
    Detects mergefield-like placeholders and IF fields in a Word document and converts them into actual merge fields and IF fields.
    Specifically handles lines like:
    { IF "J" "{ MERGEFIELD ab-borger-enlig-ved-aeldrecheck-berettigelse }" "dine" "din og din samlever/ægtefælles" }

    Args:
        docx_path (str): Path to the input .docx file.
        output_path (str, optional): Path to save the modified .docx file. If None, overwrites the input file.
    
    Returns:
        tuple: (bool, str) - Success flag and debug info
    """
    try:
        import re
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        doc = Document(docx_path)
        debug_info = []
        conversion_count = 0

        # Simplified patterns for direct match
        mergefield_pattern = re.compile(r"\{ *MERGEFIELD +([^\} ]+) *\}")
        
        # Exact pattern for the example
        if_pattern = re.compile(r'\{\s*IF\s+"J"\s+"{\s*MERGEFIELD\s+([\w\-:]+)\s*}"\s+"([^"]*)"\s+"([^"]*)"\s*\}')
        
        # More generic pattern as fallback
        if_pattern_generic = re.compile(r'\{\s*IF\s+"([^"]+)"\s+"{\s*MERGEFIELD\s+([\w\-:]+)\s*}"\s+"([^"]*)"\s+"([^"]*)"\s*\}')

        debug_info.append(f"Processing document: {docx_path}")

        # Define helper function to create merge fields
        def create_merge_field(run, field_name):
            run.text = ""
            
            # Create field begin
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            # Create instruction text
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # Important to preserve whitespace
            instrText.text = f' MERGEFIELD {field_name} '
            
            # Create field separator
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            
            # Create field end
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            # Append elements
            r = run._r
            r.append(fldChar1)
            r.append(instrText)
            r.append(fldChar2)
            r.append(fldChar3)

        # Define helper function to create IF fields
        def create_if_field(paragraph, cond, mergefield, true_text, false_text):
            # Remove all runs in the paragraph
            for run in list(paragraph.runs):
                paragraph._element.remove(run._element)
            
            # Add new run for the field
            run = paragraph.add_run()
            r = run._r
            
            # Create field begin
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            
            # Create instruction text
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # Important to preserve whitespace
            
            # Special format for the case where condition is "J"
            if cond == "J":
                instrText.text = f' IF "J" = "{{ MERGEFIELD {mergefield} }}" "{true_text}" "{false_text}" '
            else:
                instrText.text = f' IF "{cond}" "{{ MERGEFIELD {mergefield} }}" "{true_text}" "{false_text}" '
            
            # Create field separator
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            
            # Create field end
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            
            # Append elements
            r.append(fldChar1)
            r.append(instrText)
            r.append(fldChar2)
            r.append(fldChar3)

        # Process each paragraph
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            debug_info.append(f"Processing paragraph {i+1}: '{para_text[:50]}...'")
            
            # Look for the exact match first - this is crucial for the example line
            m = if_pattern.search(para_text)
            if m:
                mergefield_name = m.group(1)
                true_text = m.group(2)
                false_text = m.group(3)
                debug_info.append(f"Found exact IF pattern with mergefield: {mergefield_name}")
                create_if_field(para, "J", mergefield_name, true_text, false_text)
                conversion_count += 1
                continue
                
            # Try the generic pattern as fallback
            m = if_pattern_generic.search(para_text)
            if m:
                cond = m.group(1)
                mergefield_name = m.group(2)
                true_text = m.group(3)
                false_text = m.group(4)
                debug_info.append(f"Found generic IF pattern with condition: {cond}, mergefield: {mergefield_name}")
                create_if_field(para, cond, mergefield_name, true_text, false_text)
                conversion_count += 1
                continue
                
            # Otherwise, search for MERGEFIELDs in runs
            for run in para.runs:
                matches = list(mergefield_pattern.finditer(run.text))
                if not matches:
                    continue
                new_text = run.text
                for match in reversed(matches):
                    field_name = match.group(1)
                    before = new_text[: match.start()]
                    after = new_text[match.end() :]
                    debug_info.append(f"Found MERGEFIELD: {field_name}")
                    run.text = before
                    merge_run = para.add_run()
                    create_merge_field(merge_run, field_name)
                    if after:
                        para.add_run(after)
                    new_text = before
                    conversion_count += 1

        if output_path is None:
            output_path = docx_path
        doc.save(output_path)
        
        debug_message = f"Converted {conversion_count} fields"
        if debug_info:
            debug_message += f"\nDebug info: {', '.join(debug_info)}"
        return (conversion_count > 0, debug_message)
        
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        return (False, f"Error: {str(e)}\n{error_details}")
