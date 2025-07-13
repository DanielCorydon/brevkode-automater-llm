from docx import Document
from docx.oxml import OxmlElement
from .merge_fields import (
    create_merge_field,
    create_if_field,
    create_merge_field_with_formatting,
    create_if_field_with_formatting,
)
import re


def _get_if_field_code(condition_title, mappings):
    condition_key = mappings.get(condition_title)
    special_html_mappings = {
        "ab-ubegraenset-fuldmagt": "Html:x-fuldmagtsbetingelse",
    }
    true_result_key = None
    if condition_key and condition_key in special_html_mappings:
        true_result_key = special_html_mappings[condition_key]
    elif condition_key:
        if f"Html:{condition_key}" in mappings.values():
            for k, v in mappings.items():
                if v == f"Html:{condition_key}":
                    true_result_key = v
                    break
        if not true_result_key:
            suffix = None
            if "-" in condition_key:
                suffix = condition_key.split("-", 1)[1]
            elif ":" in condition_key:
                suffix = condition_key.split(":", 1)[1]
            if suffix:
                for k, v in mappings.items():
                    if v.startswith("Html:") and suffix in v:
                        true_result_key = v
                        break
    if (
        not true_result_key
        and condition_key
        and condition_title == "Ubegrænset fuldmagt"
    ):
        true_result_key = "Html:x-fuldmagtsbetingelse"
    return condition_key, true_result_key


def process_paragraph(paragraph, text, mappings):
    # Replace all 'If betingelse <condition>' with IF fields, rest with merge fields
    sorted_mappings = sorted(mappings.items(), key=lambda x: len(x[0]), reverse=True)
    p = paragraph._p
    for child in list(p):
        p.remove(child)

    # Regex to find all 'If betingelse <condition>' (case-insensitive)
    pattern = re.compile(r"[Ii]f betingelse ([^\”\"]+)")
    pos = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        # Add text before the match
        if start > pos:
            _add_mergefields_to_text(paragraph, text[pos:start], sorted_mappings)
        condition_title = match.group(1).strip()
        condition_key, true_result_key = _get_if_field_code(condition_title, mappings)
        run = paragraph.add_run()
        r = run._r
        if condition_key and true_result_key:
            create_if_field(r, condition_key, true_result_key)
        else:
            run.add_text(match.group(0))
        pos = end
    # Add any remaining text after last match
    if pos < len(text):
        _add_mergefields_to_text(paragraph, text[pos:], sorted_mappings)
    return paragraph


def _add_mergefields_to_text(paragraph, text, sorted_mappings):
    # Replace all mapped titles with merge fields, rest as text
    remaining_text = text
    while remaining_text:
        match_found = False
        for titel, nogle in sorted_mappings:
            idx = remaining_text.find(titel)
            if idx != -1:
                if idx > 0:
                    paragraph.add_run(remaining_text[:idx])
                run = paragraph.add_run()
                r = run._r
                create_merge_field(r, nogle)
                remaining_text = remaining_text[idx + len(titel) :]
                match_found = True
                break
        if not match_found:
            paragraph.add_run(remaining_text)
            break


def create_document_with_merge_fields(template_text, mappings):
    import re

    doc = Document()
    paragraphs = template_text.split("\n")
    debug_output = []
    pattern = re.compile(r"[Ii]f betingelse ([^\”\"]+)")
    for para_text in paragraphs:
        if para_text.strip():
            debug_line = ""
            pos = 0
            for match in pattern.finditer(para_text):
                start, end = match.span()
                # Add text before the match, replacing mapped titles
                if start > pos:
                    debug_line += _replace_titles_with_mergefields(
                        para_text[pos:start], mappings
                    )
                condition_title = match.group(1).strip()
                condition_key, true_result_key = _get_if_field_code(
                    condition_title, mappings
                )
                if condition_key and true_result_key:
                    debug_line += f'{{ IF "{{ MERGEFIELD {condition_key} }}" = "J" "din" "jeres" }}'
                else:
                    debug_line += match.group(0)
                pos = end
            # Add any remaining text after last match
            if pos < len(para_text):
                debug_line += _replace_titles_with_mergefields(
                    para_text[pos:], mappings
                )
            debug_output.append(debug_line)
            p = doc.add_paragraph()
            process_paragraph(p, para_text, mappings)
    return doc, "\n".join(debug_output)


def _replace_titles_with_mergefields(text, mappings):
    # Replace all mapped titles with merge field code strings
    sorted_mappings = sorted(mappings.items(), key=lambda x: len(x[0]), reverse=True)
    result = text
    for titel, nogle in sorted_mappings:
        result = result.replace(titel, f"{{ MERGEFIELD {nogle} }}")
    return result


def process_docx_template(doc, mappings):
    # Remove all comments from the document
    if hasattr(doc, "part") and hasattr(doc.part, "package"):
        # Remove comments part if it exists
        comments_part = None
        for rel in doc.part.rels.values():
            if (
                rel.reltype
                == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
            ):
                comments_part = rel.target_part
                break
        if comments_part:
            # Remove all comment references in the document
            for para in doc.paragraphs:
                p = para._p
                comment_refs = p.findall(".//w:commentRangeStart", namespaces=p.nsmap)
                for ref in comment_refs:
                    ref.getparent().remove(ref)
                comment_refs = p.findall(".//w:commentRangeEnd", namespaces=p.nsmap)
                for ref in comment_refs:
                    ref.getparent().remove(ref)
                comment_refs = p.findall(".//w:commentReference", namespaces=p.nsmap)
                for ref in comment_refs:
                    ref.getparent().remove(ref)
            # Remove the comments part from the package
            if comments_part.partname in doc.part.package.parts:
                doc.part.package.parts.pop(comments_part.partname)
            # Remove the relationship
            for rel_id, rel in list(doc.part.rels.items()):
                if rel.target_part == comments_part:
                    doc.part.rels.pop(rel_id)
    sorted_mappings = sorted(mappings.items(), key=lambda x: len(x[0]), reverse=True)
    debug_output = []
    for para in doc.paragraphs:
        para_text = para.text
        if not para_text.strip():
            continue
        if para_text.strip().startswith("IF Betingelse "):
            condition_title = para_text.strip()[len("IF Betingelse ") :]
            condition_key = mappings.get(condition_title)
            special_html_mappings = {
                "ab-ubegraenset-fuldmagt": "Html:x-fuldmagtsbetingelse"
            }
            true_result_key = None
            if condition_key and condition_key in special_html_mappings:
                true_result_key = special_html_mappings[condition_key]
            elif condition_key:
                if f"Html:{condition_key}" in mappings.values():
                    for k, v in mappings.items():
                        if v == f"Html:{condition_key}":
                            true_result_key = v
                            break
                if not true_result_key:
                    suffix = None
                    if "-" in condition_key:
                        suffix = condition_key.split("-", 1)[1]
                    elif ":" in condition_key:
                        suffix = condition_key.split(":", 1)[1]
                    if suffix:
                        for k, v in mappings.items():
                            if v.startswith("Html:") and suffix in v:
                                true_result_key = v
                                break
            if (
                not true_result_key
                and condition_key
                and condition_title == "Ubegrænset fuldmagt"
            ):
                true_result_key = "Html:x-fuldmagtsbetingelse"
            if condition_key and true_result_key:
                debug_para = f'{{ IF "J" = "{{ MERGEFIELD {condition_key} }}" "{{ MERGEFIELD {true_result_key} }}" }}'
            else:
                debug_para = para_text
            debug_output.append(debug_para)
            if len(para.runs) > 0:
                original_run = para.runs[0]
                for run in list(para.runs):
                    run._element.getparent().remove(run._element)
                new_run = para.add_run()
                new_run.bold = original_run.bold
                new_run.italic = original_run.italic
                new_run.underline = original_run.underline
                if original_run.font.name:
                    new_run.font.name = original_run.font.name
                if original_run.font.size:
                    new_run.font.size = original_run.font.size
                if original_run.font.color.rgb:
                    new_run.font.color.rgb = original_run.font.color.rgb
                if condition_key and true_result_key:
                    create_if_field_with_formatting(
                        new_run._element, condition_key, true_result_key, original_run
                    )
                else:
                    new_run.text = para_text
            else:
                run = para.add_run()
                if condition_key and true_result_key:
                    create_if_field_with_formatting(
                        run._element, condition_key, true_result_key
                    )
                else:
                    run.text = para_text
        else:
            run_data = [(run, run.text) for run in para.runs]
            debug_para = para_text
            for titel, nogle in sorted_mappings:
                if titel in debug_para:
                    debug_para = debug_para.replace(titel, f"{{ MERGEFIELD {nogle} }}")
            debug_output.append(debug_para)
            for i, (run, run_text) in enumerate(run_data):
                if not run_text.strip():
                    continue
                replaced = False
                for titel, nogle in sorted_mappings:
                    if titel in run_text:
                        original_format = run._element
                        parts = run_text.split(titel)
                        run._element.clear_content()
                        if parts[0]:
                            run.text = parts[0]
                        # Insert merge field and remove color/highlight from this run
                        create_merge_field_with_formatting(run._element, nogle, run)
                        run.font.color.rgb = None
                        if hasattr(run.font, "highlight_color"):
                            run.font.highlight_color = None
                        if len(parts) > 1 and parts[1]:
                            new_run = para.add_run(parts[1])
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if run.font.name:
                                new_run.font.name = run.font.name
                            if run.font.size:
                                new_run.font.size = run.font.size
                            if run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                            if (
                                hasattr(new_run.font, "highlight_color")
                                and run.font.highlight_color
                            ):
                                new_run.font.highlight_color = run.font.highlight_color
                        replaced = True
                        break
                if not replaced:
                    pass
    return doc, "\n".join(debug_output)
